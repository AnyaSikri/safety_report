"""
Template Populator Module - Stage 3
Populates the DSR Word template with matched content
"""

import re
from typing import Dict, List, Set
from pathlib import Path
from docx import Document


class TemplatePopulator:
    """Populates DSR Word template with extracted content"""
    
    def __init__(self, template_path: str):
        """
        Initialize with path to DSR Word template
        
        Args:
            template_path: Path to the DSR template .docx file
        """
        self.template_path = Path(template_path)
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        self.doc = Document(template_path)
        self.placeholders_found = set()
    
    def find_all_placeholders(self) -> List[str]:
        """
        Find all placeholders in the document
        Search in paragraphs, tables, headers, and footers
        
        Returns:
            Sorted list of unique placeholder strings
        """
        print("Scanning document for placeholders...")
        placeholders = set()
        
        # Search paragraphs
        for para in self.doc.paragraphs:
            placeholders.update(self._extract_placeholders(para.text))
        
        # Search tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        placeholders.update(self._extract_placeholders(para.text))
        
        # Search headers
        for section in self.doc.sections:
            header = section.header
            for para in header.paragraphs:
                placeholders.update(self._extract_placeholders(para.text))
            
            # Search footer
            footer = section.footer
            for para in footer.paragraphs:
                placeholders.update(self._extract_placeholders(para.text))
        
        self.placeholders_found = placeholders
        print(f"✓ Found {len(placeholders)} unique placeholders")
        
        return sorted(placeholders)
    
    def _extract_placeholders(self, text: str) -> Set[str]:
        """
        Extract [INSERT_*] placeholders from text
        
        Args:
            text: Text to search for placeholders
            
        Returns:
            Set of placeholder strings found
        """
        return set(re.findall(r'\[INSERT_[A-Z0-9_]+\]', text))
    
    def replace_placeholder(self, placeholder: str, new_content: str):
        """
        Replace a placeholder with new content throughout the document
        
        Args:
            placeholder: Placeholder string to replace (e.g., "[INSERT_DRUG_NAME]")
            new_content: Content to insert
        """
        if not new_content:
            new_content = "[NO CONTENT]"
        
        # Clean up content - remove excessive whitespace but preserve single line breaks
        new_content = self._clean_content(new_content)
        
        replacements_made = 0
        
        # Replace in paragraphs
        for para in self.doc.paragraphs:
            if placeholder in para.text:
                replacements_made += self._replace_in_paragraph(para, placeholder, new_content)
        
        # Replace in tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder in para.text:
                            replacements_made += self._replace_in_paragraph(para, placeholder, new_content)
        
        # Replace in headers
        for section in self.doc.sections:
            header = section.header
            for para in header.paragraphs:
                if placeholder in para.text:
                    replacements_made += self._replace_in_paragraph(para, placeholder, new_content)
            
            footer = section.footer
            for para in footer.paragraphs:
                if placeholder in para.text:
                    replacements_made += self._replace_in_paragraph(para, placeholder, new_content)
        
        return replacements_made
    
    def _clean_content(self, content: str) -> str:
        """
        Clean content before insertion
        
        Args:
            content: Raw content string
            
        Returns:
            Cleaned content string
        """
        if not content:
            return ""
        
        # Remove excessive whitespace (multiple spaces/tabs)
        content = re.sub(r'[ \t]+', ' ', content)
        
        # Normalize line breaks - convert multiple newlines to double newline (paragraph break)
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        # Remove trailing/leading whitespace
        content = content.strip()
        
        return content
    
    def _replace_in_paragraph(self, paragraph, old_text: str, new_text: str) -> int:
        """
        Replace text in paragraph while clearing placeholder formatting
        
        Args:
            paragraph: python-docx Paragraph object
            old_text: Text to replace
            new_text: Replacement text
            
        Returns:
            Number of replacements made (0 or 1)
        """
        if old_text not in paragraph.text:
            return 0
        
        # Build a list of run texts to find where placeholder is
        runs = paragraph.runs
        if not runs:
            return 0
        
        # Try to find which run(s) contain the placeholder
        for i, run in enumerate(runs):
            if old_text in run.text:
                # Placeholder is entirely in this run - simple case
                run.text = run.text.replace(old_text, new_text, 1)
                # Clear excessive formatting from the run
                run.bold = False
                run.font.size = None  # Reset to default size
                return 1
        
        # Placeholder might be split across runs
        # Reconstruct full text and find position
        full_text = ''.join(run.text for run in runs)
        
        if old_text not in full_text:
            return 0
        
        # Find the placeholder position
        placeholder_start = full_text.index(old_text)
        placeholder_end = placeholder_start + len(old_text)
        
        # Calculate new text
        new_full_text = full_text[:placeholder_start] + new_text + full_text[placeholder_end:]
        
        # Find which runs the placeholder spans
        current_pos = 0
        start_run_idx = -1
        end_run_idx = -1
        
        for i, run in enumerate(runs):
            run_len = len(run.text)
            if current_pos <= placeholder_start < current_pos + run_len:
                start_run_idx = i
            if current_pos < placeholder_end <= current_pos + run_len:
                end_run_idx = i
            current_pos += run_len
        
        if start_run_idx == -1 or end_run_idx == -1:
            # Fallback: clear all and replace
            paragraph.clear()
            new_run = paragraph.add_run(new_full_text)
            # Ensure no excessive formatting
            new_run.bold = False
            new_run.font.size = None  # Reset to default size
            return 1
        
        # Replace the text while clearing placeholder formatting
        if start_run_idx == end_run_idx:
            # Placeholder in single run
            run = runs[start_run_idx]
            run.text = run.text.replace(old_text, new_text, 1)
            # Clear excessive formatting
            run.bold = False
            run.font.size = None  # Reset to default size
        else:
            # Placeholder spans multiple runs - merge and replace
            # Clear formatting on the run that will contain the replacement
            position_in_text = 0
            for i in range(len(runs)):
                run_len = len(runs[i].text)
                if i < start_run_idx:
                    position_in_text += run_len
                elif i == start_run_idx:
                    # This run starts before or at placeholder
                    start_offset = placeholder_start - position_in_text
                    runs[i].text = runs[i].text[:start_offset] + new_text
                    # Clear excessive formatting
                    runs[i].bold = False
                    runs[i].font.size = None  # Reset to default size
                    position_in_text += run_len
                elif start_run_idx < i < end_run_idx:
                    # This run is completely within placeholder - clear it
                    runs[i].text = ''
                    position_in_text += run_len
                elif i == end_run_idx:
                    # This run contains end of placeholder
                    end_offset = placeholder_end - position_in_text
                    runs[i].text = runs[i].text[end_offset:]
                    position_in_text += run_len
        
        return 1
    
    def populate_table_row(self, table_index: int, row_index: int, column_values: List[str]):
        """
        Populate a specific table row with values
        
        Args:
            table_index: Index of the table in the document
            row_index: Index of the row in the table
            column_values: List of values for each column
        """
        if table_index >= len(self.doc.tables):
            print(f"⚠ Warning: Table index {table_index} out of range")
            return
        
        table = self.doc.tables[table_index]
        
        if row_index >= len(table.rows):
            print(f"⚠ Warning: Row index {row_index} out of range")
            return
        
        row = table.rows[row_index]
        
        for col_index, value in enumerate(column_values):
            if col_index < len(row.cells):
                row.cells[col_index].text = str(value)
    
    def populate_all_fields(self, field_content_dict: Dict[str, str]):
        """
        Populate all placeholders with matched content
        
        Args:
            field_content_dict: Dict mapping placeholders to content
        """
        print("\n" + "="*60)
        print("Populating DSR Template")
        print("="*60)
        
        # Find all placeholders if not already done
        if not self.placeholders_found:
            self.find_all_placeholders()
        
        populated = 0
        skipped = 0
        
        for placeholder, content in field_content_dict.items():
            # Skip empty or error content
            if not content or content == "N/A":
                print(f"⊘ Skipping {placeholder}: No content")
                skipped += 1
                continue
            
            # Check if placeholder exists in document
            if placeholder not in self.placeholders_found:
                print(f"⚠ Warning: {placeholder} not found in template")
                skipped += 1
                continue
            
            # Replace the placeholder
            replacements = self.replace_placeholder(placeholder, content)
            
            if replacements > 0:
                print(f"✓ Replaced {placeholder} ({len(content)} chars, {replacements} locations)")
                populated += 1
            else:
                print(f"⊘ Could not replace {placeholder}")
                skipped += 1
        
        print("\n" + "="*60)
        print(f"Population Summary: {populated} populated, {skipped} skipped")
        print("="*60 + "\n")
    
    def save_document(self, output_path: str):
        """
        Save the populated document
        
        Args:
            output_path: Path where to save the populated document
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        self.doc.save(str(output_path))
        print(f"✓ Document saved to: {output_path}")
    
    def generate_mapping_report(self, field_content_dict: Dict[str, str]) -> Dict[str, List[str]]:
        """
        Generate a report showing which fields were populated
        
        Args:
            field_content_dict: Dict mapping placeholders to content
            
        Returns:
            Dict with categorized field lists
        """
        report = {
            'populated': [],
            'empty': [],
            'not_in_ib': [],
            'errors': []
        }
        
        for placeholder, content in field_content_dict.items():
            if not content or content == "N/A" or len(content.strip()) == 0:
                report['empty'].append(placeholder)
            elif "NOT AVAILABLE IN IB" in content or "REQUIRES:" in content:
                report['not_in_ib'].append(placeholder)
            elif "[ERROR" in content or "error" in content.lower():
                report['errors'].append(placeholder)
            else:
                report['populated'].append(placeholder)
        
        return report
    
    def print_report(self, report: Dict[str, List[str]]):
        """
        Print a formatted report
        
        Args:
            report: Report dict from generate_mapping_report
        """
        print("\n" + "="*60)
        print("POPULATION REPORT")
        print("="*60)
        
        print(f"\n✓ Successfully Populated: {len(report['populated'])}")
        if report['populated']:
            for field in sorted(report['populated'])[:10]:  # Show first 10
                print(f"  - {field}")
            if len(report['populated']) > 10:
                print(f"  ... and {len(report['populated']) - 10} more")
        
        print(f"\n⊘ Empty/No Content: {len(report['empty'])}")
        if report['empty']:
            for field in sorted(report['empty'])[:5]:
                print(f"  - {field}")
            if len(report['empty']) > 5:
                print(f"  ... and {len(report['empty']) - 5} more")
        
        print(f"\n⚠ Not Available in IB: {len(report['not_in_ib'])}")
        if report['not_in_ib']:
            for field in sorted(report['not_in_ib'])[:5]:
                print(f"  - {field}")
            if len(report['not_in_ib']) > 5:
                print(f"  ... and {len(report['not_in_ib']) - 5} more")
        
        print(f"\n✗ Errors: {len(report['errors'])}")
        if report['errors']:
            for field in sorted(report['errors']):
                print(f"  - {field}")
        
        print("\n" + "="*60)


def main():
    """Standalone execution for testing"""
    import argparse
    import json
    
    parser = argparse.ArgumentParser(description='Populate DSR template')
    parser.add_argument('--template', required=True, help='Path to DSR template')
    parser.add_argument('--content', required=True, help='Path to matched content JSON')
    parser.add_argument('--output', required=True, help='Path for output document')
    
    args = parser.parse_args()
    
    # Load matched content
    with open(args.content, 'r', encoding='utf-8') as f:
        content_dict = json.load(f)
    
    # Populate template
    populator = TemplatePopulator(args.template)
    populator.populate_all_fields(content_dict)
    
    # Generate and print report
    report = populator.generate_mapping_report(content_dict)
    populator.print_report(report)
    
    # Save document
    populator.save_document(args.output)


if __name__ == '__main__':
    main()

